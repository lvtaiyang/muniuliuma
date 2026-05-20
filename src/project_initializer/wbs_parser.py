"""WBS 解析器 — LLM 从合同/清单/分部分项文档中提取结构化数据。"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from openai import OpenAI
from .. import config as cfg

PARSE_SYSTEM_PROMPT = """你是工程建设领域的项目结构分析专家。从用户提供的项目文档中提取合同信息和分部分项结构。

## 你需要提取

### 1. 合同关键信息
- 合同名称、编号
- 发包人（甲方）、承包人（乙方）
- 监理单位
- 合同金额
- 计划开工/竣工日期
- 质量等级要求
- 其他关键条款摘要

### 2. 分部分项结构（WBS）
按照 单位工程 → 分部工程 → 分项工程 → 检验批 的层级：
- 单位工程：可独立发挥功能的工程（如 1栋楼、1段路基）
- 分部工程：按专业/部位划分（如 地基与基础、主体结构）
- 分项工程：按工种/工序划分（如 土方开挖、钢筋工程）
- 检验批：同一条件下施工的批次

每层包含：层级、编号、名称、类型

## 返回格式

{
  "contract": {
    "contract_name": "",
    "contract_no": "",
    "party_a": "",
    "party_b": "",
    "supervisor": "",
    "contract_amount": "",
    "start_date": "",
    "end_date": "",
    "quality_grade": "合格",
    "key_clauses": ["关键条款1", "条款2"]
  },
  "wbs": [
    {
      "level": 1,
      "code": "01",
      "name": "单位工程或分部名称",
      "type": "单位工程|分部工程|分项工程|检验批",
      "children": [
        {
          "level": 2,
          "code": "01-01",
          "name": "子项名称",
          "type": "分项工程",
          "children": []
        }
      ]
    }
  ],
  "notes": "解析说明，如数据来源、缺失信息等"
}

## 规则
- 如果用户只给了分部分项，contract 字段可为空对象
- 如果文档是 xlsx，注意表头行可能在第1-3行
- 分部分项的编号按文档原样保留，不要编造
- 缺失的信息用空字符串，不要猜测
- 检验批如果文档没明确列出来，可以省略
"""


def parse_project_document(file_path: str | Path) -> dict[str, Any]:
    """从项目文档中提取合同和分部分项信息。

    支持: .xlsx .docx .csv .txt .md .yaml .json
    """
    file_path = Path(file_path)
    if not file_path.exists():
        return {"error": f"文件不存在: {file_path}"}

    text = _read_document(file_path)
    if not text or text.startswith("[无法"):
        return {"error": f"无法读取文件: {text}"}

    llm = cfg.get_llm_config("text")
    if not llm.get("api_key"):
        return {"error": "未配置 LLM text.api_key"}

    client = OpenAI(base_url=llm["base_url"], api_key=llm["api_key"], timeout=180)

    response = client.chat.completions.create(
        model=llm["model"],
        messages=[
            {"role": "system", "content": PARSE_SYSTEM_PROMPT},
            {"role": "user", "content": (
                f"文件名: {file_path.name}\n\n"
                f"文档内容:\n{text[:15000]}\n\n"
                f"请提取合同信息和分部分项结构。"
            )},
        ],
        max_tokens=4000,
        temperature=0.1,
    )

    return _parse_response(response.choices[0].message.content or "")


def parse_from_text(text: str, context: str = "") -> dict[str, Any]:
    """从用户提供的文本中解析项目结构。适用场景：用户直接描述分部分项。"""
    llm = cfg.get_llm_config("text")
    if not llm.get("api_key"):
        return {"error": "未配置 LLM text.api_key"}

    client = OpenAI(base_url=llm["base_url"], api_key=llm["api_key"], timeout=120)

    ctx = f"\n\n补充说明: {context}" if context else ""

    response = client.chat.completions.create(
        model=llm["model"],
        messages=[
            {"role": "system", "content": PARSE_SYSTEM_PROMPT},
            {"role": "user", "content": (
                f"用户提供的项目信息:\n{text[:8000]}{ctx}\n\n"
                f"请提取合同信息和分部分项结构。如果信息不完整，已提供的尽量提取，缺失的留空。"
            )},
        ],
        max_tokens=4000,
        temperature=0.1,
    )
    return _parse_response(response.choices[0].message.content or "")


def _read_document(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    try:
        if ext in (".txt", ".md", ".markdown"):
            return file_path.read_text(encoding="utf-8", errors="replace")[:15000]
        if ext == ".csv":
            content = file_path.read_text(encoding="utf-8", errors="replace")
            return f"CSV:\n{content[:15000]}"
        if ext in (".xlsx", ".xlsm"):
            return _read_xlsx_text(file_path)
        if ext == ".docx":
            return _read_docx_text(file_path)
        if ext in (".yaml", ".yml", ".json"):
            content = file_path.read_text(encoding="utf-8", errors="replace")
            return content[:15000]
    except Exception as e:
        return f"[读取失败: {e}]"
    return f"[不支持的文件格式: {ext}]"


def _read_xlsx_text(file_path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        lines = []
        for sn in wb.sheetnames:
            ws = wb[sn]
            lines.append(f"\n=== 工作表: {sn} ===")
            for row in ws.iter_rows(values_only=True, max_row=300):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(lines)[:15000]
    except ImportError:
        return "[需要 openpyxl: pip install openpyxl]"


def _read_docx_text(file_path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(file_path))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)[:15000]
    except ImportError:
        return "[需要 python-docx: pip install python-docx]"


def _parse_response(content: str) -> dict[str, Any]:
    from .. import json_utils
    result = json_utils.parse_llm_json(content)
    if "_raw" in result:
        return {"raw": content, "contract": {}, "wbs": [],
                "notes": "LLM 返回格式异常，请检查原始输出"}
    return result
