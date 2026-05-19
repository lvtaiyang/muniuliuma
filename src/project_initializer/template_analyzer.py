"""模板分析：用 LLM 分析用户上传的文档模板（Word/Excel/PDF），
提取字段结构、使用场景，生成结构化模板定义。"""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from openai import OpenAI

SYSTEM_PROMPT = """你是工程建设领域的文档分析专家。分析用户上传的文档/表格，提取模板结构。

你必须严格返回 JSON 格式，不要包含任何其他文字：

{
  "template_name": "模板名称（简短准确，如'材料进场检验表'）",
  "description": "模板用途描述（一句话）",
  "applies_to": {
    "stage": "适用的建设阶段代码，从以下选：01_前期决策, 02_设计准备, 03_招标合同, 04_施工实施, 05_竣工验收, 06_结算审计, 07_后评估",
    "scenario": "具体使用场景，如'材料进场验收''隐蔽工程验收'"
  },
  "fields": [
    {
      "name": "字段名",
      "type": "string|number|date|enum|text|image|file",
      "required": true,
      "description": "字段含义说明",
      "options": ["仅enum类型填写，如: 合格,不合格"]
    }
  ],
  "notes": "使用注意事项或填写说明"
}

字段类型说明：
- string: 短文本（名称、编号等）
- number: 数值（数量、金额等）
- date: 日期
- enum: 枚举选项
- text: 长文本（备注、说明等）
- image: 图片/照片
- file: 附件文件
"""


def analyze_document(
    file_path: str | Path,
    base_url: str,
    api_key: str,
    model: str,
) -> dict[str, Any]:
    """分析用户上传的文档模板，提取结构化定义。

    支持 .docx .xlsx .pdf 格式。
    先读取文件内容，再送 LLM 提取字段。
    """
    file_path = Path(file_path)
    if not file_path.exists():
        return {"error": f"文件不存在: {file_path}"}

    text_content = _read_file(file_path)

    client = OpenAI(base_url=base_url, api_key=api_key, timeout=120)

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {
                "role": "user",
                "content": (
                    f"文件名: {file_path.name}\n\n"
                    f"文件内容:\n{text_content}\n\n"
                    f"请分析这个模板，提取字段结构和适用场景。"
                ),
            },
        ],
        max_tokens=2000,
        temperature=0.1,
    )

    content = response.choices[0].message.content or ""
    return _parse_response(content)


def _read_file(file_path: Path) -> str:
    """读取文档内容为文本。"""
    ext = file_path.suffix.lower()

    if ext == ".txt":
        return file_path.read_text(encoding="utf-8", errors="replace")

    if ext in (".md", ".markdown"):
        return file_path.read_text(encoding="utf-8", errors="replace")

    if ext in (".csv",):
        content = file_path.read_text(encoding="utf-8", errors="replace")
        return f"CSV 文件内容（前 200 行）:\n{chr(10).join(content.splitlines()[:200])}"

    if ext in (".docx",):
        try:
            return _read_docx(file_path)
        except ImportError:
            return f"[无法解析 .docx 文件，请安装 python-docx: pip install python-docx]\n文件名: {file_path.name}"

    if ext in (".xlsx", ".xlsm"):
        try:
            return _read_xlsx(file_path)
        except ImportError:
            return f"[无法解析 .xlsx 文件，请安装 openpyxl: pip install openpyxl]\n文件名: {file_path.name}"

    if ext == ".pdf":
        try:
            return _read_pdf(file_path)
        except ImportError:
            return f"[无法解析 .pdf 文件]\n文件名: {file_path.name}"

    # 兜底：只给文件名让 LLM 从名称推断
    return f"[无法解析的文件格式: {ext}]\n文件名: {file_path.name}\n请根据文件名推断模板结构。"


def _read_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]

    # 也读取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))

    return "\n".join(lines[:300])


def _read_xlsx(path: Path) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"\n=== 工作表: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True, max_row=200):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_pdf(path: Path) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            lines = []
            for page in pdf.pages[:10]:
                text = page.extract_text()
                if text:
                    lines.append(text)
            return "\n".join(lines)
    except ImportError:
        pass
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(str(path))
        lines = []
        for page in reader.pages[:10]:
            text = page.extract_text()
            if text:
                lines.append(text)
        return "\n".join(lines)
    except ImportError:
        raise ImportError("请安装 pdfplumber 或 PyPDF2: pip install pdfplumber")


def _parse_response(content: str) -> dict[str, Any]:
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        pass

    match = re.search(r"```(?:json)?\s*\n?([\s\S]*?)\n?```", content)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    match = re.search(r"\{[\s\S]*\}", content)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass

    return {"raw_response": content, "template_name": "无法识别",
            "description": "", "fields": [], "applies_to": {}}
