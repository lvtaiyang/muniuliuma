"""COM 自动化助手 — 通过 win32com 驱动本地 Word/Excel/WPS 操作文档。

优先 win32com（完美保留格式），无 COM 环境时自动回退到 openpyxl/python-docx。
"""

from __future__ import annotations

import re
import time
from contextlib import contextmanager
from pathlib import Path
from typing import Any

# win32com 仅在 Windows 上可用
try:
    import win32com.client as win32
    from win32com.client import Dispatch, constants
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False


def _check_openpyxl():
    try:
        import openpyxl
        return True
    except ImportError:
        return False


def _check_python_docx():
    try:
        import docx
        return True
    except ImportError:
        return False


# ── 应用探测 ────────────────────────────────────────────────────

def _detect_word_app():
    if not WIN32_AVAILABLE:
        return None, ""
    candidates = [
        ("Kwps.Application", "WPS Word"),
        ("Word.Application", "Microsoft Word"),
    ]
    for prog_id, name in candidates:
        try:
            app = Dispatch(prog_id)
            app.Visible = False
            return app, name
        except Exception:
            continue
    return None, ""


def _detect_excel_app():
    if not WIN32_AVAILABLE:
        return None, ""
    candidates = [
        ("Ket.Application", "WPS Excel"),
        ("Excel.Application", "Microsoft Excel"),
    ]
    for prog_id, name in candidates:
        try:
            app = Dispatch(prog_id)
            app.Visible = False
            app.DisplayAlerts = False
            return app, name
        except Exception:
            continue
    return None, ""


def check_available() -> dict[str, Any]:
    """检测 COM 环境是否可用。"""
    if not WIN32_AVAILABLE:
        result = {
            "win32_available": False,
            "mode": "fallback",
            "word": {"available": _check_python_docx(), "method": "python-docx"},
            "excel": {"available": _check_openpyxl(), "method": "openpyxl"},
            "warning": "win32com 未安装，使用 python-docx/openpyxl 回退模式（格式保留不完美）",
        }
        return result

    word_app, word_name = _detect_word_app()
    excel_app, excel_name = _detect_excel_app()

    result = {
        "win32_available": True,
        "mode": "win32com",
        "word": {"available": bool(word_app), "app_name": word_name},
        "excel": {"available": bool(excel_app), "app_name": excel_name},
    }

    if word_app:
        try:
            word_app.Quit()
        except Exception:
            pass
    if excel_app:
        try:
            excel_app.Quit()
        except Exception:
            pass

    return result


# ── Word 操作 ───────────────────────────────────────────────────

def word_fill_placeholders(
    template_path: str | Path,
    output_path: str | Path,
    data: dict[str, str],
) -> Path:
    """替换 Word 模板中的 {{key}} 占位符并另存。

    Windows: 用 Word/WPS 原生 Find/Replace（完美保留格式）
    其他: 用 python-docx（基本保留格式）
    """
    template_path = Path(template_path).resolve()
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if WIN32_AVAILABLE:
        return _word_fill_win32(template_path, output_path, data)
    else:
        return _word_fill_docx(template_path, output_path, data)


def _word_fill_win32(template_path: Path, output_path: Path, data: dict[str, str]) -> Path:
    app, app_name = _detect_word_app()
    if not app:
        raise RuntimeError("未检测到 Word 或 WPS")

    app.Visible = False
    app.DisplayAlerts = False
    doc = None
    try:
        doc = app.Documents.Open(str(template_path))
        story_ranges = _get_all_story_ranges(doc)
        for story_range in story_ranges:
            find = story_range.Find
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                find.Text = placeholder
                find.Replacement.Text = str(value) if value else ""
                find.Forward = True
                find.Wrap = 1
                find.Format = False
                find.MatchCase = True
                find.MatchWholeWord = False
                try:
                    find.Execute(Replace=2)  # wdReplaceAll
                except Exception:
                    pass
        doc.SaveAs2(str(output_path))
        return output_path
    finally:
        try:
            if doc:
                doc.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            app.Quit()
        except Exception:
            pass


def _word_fill_docx(template_path: Path, output_path: Path, data: dict[str, str]) -> Path:
    import docx
    doc = docx.Document(str(template_path))
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.text:
                continue
            replaced = run.text
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in replaced:
                    replaced = replaced.replace(placeholder, str(value) if value else "")
            if replaced != run.text:
                run.text = replaced
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.text:
                            continue
                        replaced = run.text
                        for key, value in data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in replaced:
                                replaced = replaced.replace(placeholder, str(value) if value else "")
                        if replaced != run.text:
                            run.text = replaced
    doc.save(str(output_path))
    return output_path


def _get_all_story_ranges(doc):
    ranges = []
    try:
        ranges.append(doc.Content)
        for i in range(1, 12):
            try:
                story = doc.StoryRanges(i)
                while story:
                    ranges.append(story)
                    try:
                        story = story.NextStoryRange
                    except Exception:
                        break
            except Exception:
                continue
    except Exception:
        try:
            ranges.append(doc.Content)
        except Exception:
            pass
    return ranges


# ── Excel 操作 ──────────────────────────────────────────────────

@contextmanager
def open_excel(template_path: str | Path):
    """上下文管理器：打开 Excel 工作簿。

    Windows: 通过 COM 打开 Excel/WPS
    其他: 通过 openpyxl 打开
    """
    template_path = Path(template_path).resolve()
    if WIN32_AVAILABLE:
        app, app_name = _detect_excel_app()
        if app:
            app.Visible = False
            app.DisplayAlerts = False
            app.ScreenUpdating = False
            wb = None
            try:
                wb = app.Workbooks.Open(str(template_path))
                yield app, wb, app_name
            finally:
                try:
                    if wb:
                        wb.Close(SaveChanges=False)
                except Exception:
                    pass
                try:
                    app.Quit()
                except Exception:
                    pass
        else:
            raise RuntimeError("未检测到 Excel 或 WPS 表格")
    else:
        import openpyxl
        wb = openpyxl.load_workbook(str(template_path))
        try:
            yield None, wb, "openpyxl"
        finally:
            pass  # openpyxl 不需要显式关闭


def excel_read_structure(template_path: str | Path) -> dict[str, Any]:
    """读取 xlsx 的完整结构信息。"""
    template_path = Path(template_path).resolve()
    if not template_path.exists():
        return {"error": f"文件不存在: {template_path}"}

    if WIN32_AVAILABLE:
        return _excel_read_structure_win32(template_path)
    else:
        return _excel_read_structure_openpyxl(template_path)


def _excel_read_structure_win32(file_path: Path) -> dict[str, Any]:
    app, app_name = _detect_excel_app()
    if not app:
        raise RuntimeError("未检测到 Excel 或 WPS")

    app.Visible = False
    app.DisplayAlerts = False
    structure: dict[str, Any] = {"filename": file_path.name, "sheets": []}
    wb = None
    try:
        wb = app.Workbooks.Open(str(file_path))
        for i in range(1, wb.Worksheets.Count + 1):
            ws = wb.Worksheets(i)
            sheet_info = {"name": ws.Name, "rows": [], "merged_ranges": []}
            try:
                used = ws.UsedRange
                max_row = min(used.Rows.Count, 100)
                max_col = min(used.Columns.Count, 50)
                for row_idx in range(1, max_row + 1):
                    cells_info = []
                    for col_idx in range(1, max_col + 1):
                        cell = ws.Cells(row_idx, col_idx)
                        value = cell.Value
                        if value is not None:
                            col_letter = _col_to_letter(col_idx)
                            cells_info.append({
                                "cell": f"{col_letter}{row_idx}",
                                "value": str(value)[:200],
                                "column": col_letter,
                                "row": row_idx,
                            })
                            if cell.MergeCells:
                                merge_range = cell.MergeArea.Address.replace("$", "")
                                if merge_range not in sheet_info["merged_ranges"]:
                                    sheet_info["merged_ranges"].append(merge_range)
                    if cells_info:
                        sheet_info["rows"].append({"row": row_idx, "cells": cells_info})
            except Exception as e:
                sheet_info["error"] = str(e)
            structure["sheets"].append(sheet_info)
    finally:
        try:
            if wb:
                wb.Close(SaveChanges=False)
        except Exception:
            pass
        try:
            app.Quit()
        except Exception:
            pass
    return structure


def _excel_read_structure_openpyxl(file_path: Path) -> dict[str, Any]:
    import openpyxl
    from openpyxl.utils import get_column_letter
    wb = openpyxl.load_workbook(str(file_path), data_only=True)
    structure: dict[str, Any] = {"filename": file_path.name, "sheets": []}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_info = {
            "name": sheet_name,
            "dimensions": ws.dimensions,
            "merged_ranges": [str(m) for m in ws.merged_cells.ranges],
            "rows": [],
        }
        for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 100), values_only=False), start=1):
            cells_info = []
            for cell in row:
                if cell.value is not None:
                    col_letter = get_column_letter(cell.column)
                    cells_info.append({
                        "cell": f"{col_letter}{row_idx}",
                        "value": str(cell.value)[:200],
                        "column": col_letter,
                        "row": row_idx,
                    })
            if cells_info:
                sheet_info["rows"].append({"row": row_idx, "cells": cells_info})
        structure["sheets"].append(sheet_info)
    return structure


def excel_save_as(wb, output_path: str | Path) -> Path:
    """工作簿另存为。openpyxl 模式下直接 save，COM 模式下 SaveAs。"""
    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if WIN32_AVAILABLE:
        wb.SaveAs(str(output_path))
    else:
        wb.save(str(output_path))
    return output_path


def excel_write_cell(wb, sheet_name: str, cell_ref: str, value: Any) -> None:
    """向指定单元格写入值。"""
    if WIN32_AVAILABLE:
        ws = wb.Worksheets(sheet_name)
        ws.Range(cell_ref).Value = value
    else:
        from openpyxl.utils import column_index_from_string
        ws = wb[sheet_name] if sheet_name else wb.active
        col_str = "".join(c for c in cell_ref if c.isalpha())
        row_str = "".join(c for c in cell_ref if c.isdigit())
        if col_str and row_str:
            col_idx = column_index_from_string(col_str)
            row_idx = int(row_str)
            ws.cell(row=row_idx, column=col_idx, value=value)


def _col_to_letter(col: int) -> str:
    result = ""
    while col > 0:
        col, remainder = divmod(col - 1, 26)
        result = chr(65 + remainder) + result
    return result
