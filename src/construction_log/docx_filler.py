"""Word 模板填充：将结构化数据填入用户上传的 .docx 模板中。

模板使用 {{field_name}} 占位符，支持段落和表格。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any


def fill_template(
    template_path: str | Path,
    data: dict[str, Any],
    output_path: str | Path,
) -> Path:
    """将数据填入 Word 模板并保存。

    Args:
        template_path: 用户上传的 .docx 模板路径
        data: 键值对数据，如 {"日期": "2026-05-19", "天气": "晴", ...}
        output_path: 输出文件路径
    """
    try:
        import docx
    except ImportError:
        raise ImportError("需要 python-docx: pip install python-docx")

    doc = docx.Document(str(template_path))

    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, data)

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, data)

    # 处理页眉页脚
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, data)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _replace_in_paragraph(paragraph, data: dict[str, Any]) -> None:
    """替换段落中所有 {{key}} 占位符。保留原格式。"""
    for run in paragraph.runs:
        if not run.text:
            continue
        replaced = run.text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in replaced:
                value_str = _format_value(value)
                replaced = replaced.replace(placeholder, value_str)
        if replaced != run.text:
            run.text = replaced


def _format_value(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(f"• {v}" for v in value)
    if isinstance(value, dict):
        return "\n".join(f"{k}: {v}" for k, v in value.items())
    return str(value) if value else ""


# ── 模板占位符约定 ──────────────────────────────────────────────

# 施工日志模板推荐的占位符，用户模板中用这些即可自动填充
LOG_PLACEHOLDERS = {
    "日期": "date",
    "天气": "weather",
    "温度": "temperature",
    "施工进度": "progress",
    "材料进场": "materials",
    "质量检验": "quality",
    "安全检查": "safety",
    "关键事件": "key_events",
    "次日计划": "next_day_plan",
    "备注": "notes",
}


def build_log_data(log_json: dict[str, Any]) -> dict[str, str]:
    """将 LLM 返回的日志 JSON 转为模板填充用的扁平化键值对。"""
    data: dict[str, str] = {}

    data["日期"] = log_json.get("date", "")
    data["天气"] = log_json.get("weather", "")
    data["温度"] = log_json.get("temperature", "")

    sections = log_json.get("sections", [])
    for s in sections:
        heading = s.get("heading", "")
        content = s.get("content", "")
        # 映射常见标题到模板占位符
        key = _map_heading(heading)
        data[key] = content

    events = log_json.get("key_events", [])
    data["关键事件"] = "\n".join(f"{i+1}. {e}" for i, e in enumerate(events))

    data["次日计划"] = log_json.get("next_day_plan", "")
    data["备注"] = log_json.get("notes", "")

    return data


def _map_heading(heading: str) -> str:
    mapping = {
        "施工进度": "施工进度",
        "进度": "施工进度",
        "材料进场": "材料进场",
        "材料": "材料进场",
        "质量检验": "质量检验",
        "质量": "质量检验",
        "安全检查": "安全检查",
        "安全": "安全检查",
    }
    for pat, key in mapping.items():
        if pat in heading:
            return key
    return heading
